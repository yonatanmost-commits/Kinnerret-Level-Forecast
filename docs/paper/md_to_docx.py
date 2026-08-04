"""
md_to_docx.py - Build the Word .docx of the Kinneret paper from its markdown source.

Markdown stays the source of truth (per the project's "keep source editable, re-fit by
rebuild" principle); run this to regenerate the .docx after any edit:

    python docs/paper/md_to_docx.py

Handles the markdown subset this paper uses: # / ## / ### headings, pipe tables,
**bold** / *italic* / `code`, bullet + numbered lists, blockquotes, horizontal rules,
and LaTeX math ($inline$ and $$display$$). There is no pandoc/latex2mathml on this
machine, so math is rendered as clean Unicode (legible for a draft); a later pandoc or
LaTeX pass would give fully typeset equations.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
MD_PATH = HERE / "Kinneret_Forecast_Methods_Results.md"
DOCX_PATH = HERE / "Kinneret_Forecast_Methods_Results.docx"

_SUP = {"0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
        "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
        "-": "⁻", "+": "⁺", "n": "ⁿ", "i": "ⁱ",
        # uppercase modifier-letter superscripts (for ET0^{HS}, ^{PM}, ...)
        "A": "ᴬ", "B": "ᴮ", "D": "ᴰ", "E": "ᴱ", "G": "ᴳ", "H": "ᴴ",
        "I": "ᴵ", "J": "ᴶ", "K": "ᴷ", "L": "ᴸ", "M": "ᴹ", "N": "ᴺ",
        "O": "ᴼ", "P": "ᴾ", "R": "ᴿ", "T": "ᵀ", "U": "ᵁ", "V": "ⱽ",
        "W": "ᵂ", "S": "ˢ"}
_SUB = {"0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄",
        "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
        "-": "₋", "+": "₊", "a": "ₐ", "e": "ₑ", "i": "ᵢ",
        "j": "ⱼ", "m": "ₘ", "n": "ₙ", "o": "ₒ", "p": "ₚ",
        "s": "ₛ", "t": "ₜ", "x": "ₓ", "r": "ᵣ", "u": "ᵤ",
        "v": "ᵥ", "h": "ₕ", "k": "ₖ", "l": "ₗ"}

_GREEK = {
    r"\varphi": "φ", r"\phi": "φ", r"\delta": "δ", r"\pi": "π",
    r"\omega_s": "ωₛ", r"\omega": "ω", r"\lambda": "λ",
    r"\sigma": "σ", r"\mu": "μ", r"\tau": "τ", r"\Sigma": "Σ",
}


def _scriptify(s: str, table: dict) -> str:
    out = []
    for ch in s:
        if ch not in table:
            return None
        out.append(table[ch])
    return "".join(out)


def latex_to_unicode(tex: str) -> str:
    s = tex.strip()
    if s.startswith("$$") and s.endswith("$$"):
        s = s[2:-2]
    elif s.startswith("$") and s.endswith("$"):
        s = s[1:-1]
    s = s.strip()

    # Order matters. Word/symbol commands FIRST (so e.g. _{\max} -> _{max} before
    # subscripting), THEN scripts (so T_{max} -> Tₘₐₓ before \sqrt parses), THEN
    # brace constructs \frac/\sqrt (whose inner braces are now resolved), THEN cleanup.

    # 1. \text/\mathrm/\operatorname{...} -> inner
    s = re.sub(r"\\(?:text|mathrm|operatorname)\s*\{([^{}]*)\}", r"\1", s)
    # 2. Greek + word/symbol commands
    for k, v in _GREEK.items():
        s = s.replace(k, v)
    repl = {
        r"\cdot": "·", r"\times": "×", r"\approx": "≈",
        r"\leq": "≤", r"\le": "≤", r"\geq": "≥", r"\ge": "≥",
        r"\neq": "≠", r"\pm": "±", r"\to": "→", r"\in": "∈",
        r"\sin": "sin", r"\cos": "cos", r"\arccos": "arccos", r"\tan": "tan",
        r"\max": "max", r"\min": "min", r"\left": "", r"\right": "",
        r"\big": "", r"\Big": "", r"\qquad": "    ", r"\quad": "  ",
        r"\,": " ", r"\;": " ", r"\!": "", "~": " ",
    }
    for k, v in repl.items():
        s = s.replace(k, v)

    # 3. superscripts ^{...} / ^x  and subscripts _{...} / _x
    def _sup_brace(m):
        u = _scriptify(m.group(1), _SUP)
        return u if u is not None else "^(" + m.group(1) + ")"

    def _sub_brace(m):
        u = _scriptify(m.group(1), _SUB)
        return u if u is not None else "_(" + m.group(1) + ")"

    s = re.sub(r"\^\{([^{}]*)\}", _sup_brace, s)
    s = re.sub(r"_\{([^{}]*)\}", _sub_brace, s)
    s = re.sub(r"\^(\w)", lambda m: _SUP.get(m.group(1), "^" + m.group(1)), s)
    s = re.sub(r"_(\w)", lambda m: _SUB.get(m.group(1), "_" + m.group(1)), s)

    # 4. \frac{a}{b} -> (a)/(b); \sqrt{...} -> √(...)
    for _ in range(3):
        s = re.sub(r"\\frac\s*\{([^{}]*)\}\s*\{([^{}]*)\}", r"(\1)/(\2)", s)
    s = re.sub(r"\\sqrt\s*\{([^{}]*)\}", r"√(\1)", s)
    s = s.replace(r"\sqrt", "√")

    # 5. cleanup
    s = s.replace("{", "").replace("}", "").replace("\\", "")
    return s


# inline tokens: math, code, bold, italic  (bold before italic)
_INLINE = re.compile(
    r"(\$[^$]+\$)"          # 1 math
    r"|(`[^`]+`)"           # 2 code
    r"|(\*\*[^*]+\*\*)"     # 3 bold
    r"|(\*[^*]+\*)"         # 4 italic
)


def add_runs(p, text):
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        math, code, bold, italic = m.group(1), m.group(2), m.group(3), m.group(4)
        if math:
            r = p.add_run(latex_to_unicode(math)); r.italic = True
        elif code:
            r = p.add_run(code[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif bold:
            r = p.add_run(bold[2:-2]); r.bold = True
        elif italic:
            r = p.add_run(italic[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def is_table_sep(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c) for c in cells)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main():
    md = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(md)
    while i < n:
        line = md[i]
        stripped = line.strip()

        # blank
        if not stripped:
            i += 1
            continue
        # horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue
        # table block
        if stripped.startswith("|") and i + 1 < n and is_table_sep(md[i + 1]):
            header = split_row(md[i])
            rows = []
            j = i + 2
            while j < n and md[j].strip().startswith("|"):
                rows.append(split_row(md[j]))
                j += 1
            ncol = len(header)
            t = doc.add_table(rows=1, cols=ncol)
            try:
                t.style = "Light Grid Accent 1"
            except KeyError:
                t.style = "Table Grid"
            for c, htext in enumerate(header):
                cell = t.rows[0].cells[c]
                cell.paragraphs[0].text = ""
                add_runs(cell.paragraphs[0], htext)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = t.add_row().cells
                for c in range(ncol):
                    txt = row[c] if c < len(row) else ""
                    cells[c].paragraphs[0].text = ""
                    add_runs(cells[c].paragraphs[0], txt)
            doc.add_paragraph()
            i = j
            continue
        # headings
        m = re.match(r"(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1:
                p = doc.add_paragraph(style="Title")
                add_runs(p, text)
            else:
                p = doc.add_paragraph(style=f"Heading {level - 1}")
                add_runs(p, text)
            i += 1
            continue
        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs(p, stripped.lstrip("> ").strip())
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue
        # display math $$...$$ (single line)
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(latex_to_unicode(stripped))
            r.italic = True
            r.font.size = Pt(12)
            i += 1
            continue
        # bullet list
        if re.match(r"[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue
        # numbered list
        if re.match(r"\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue
        # normal paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")
    print(f"  paragraphs: {len(doc.paragraphs)}  tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
